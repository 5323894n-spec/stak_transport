# -*- coding: utf-8 -*-
"""Contract tests for the portrait Word route passport."""
from __future__ import annotations

import io
from dataclasses import replace
import zipfile
from datetime import date
from xml.etree import ElementTree as ET

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT

from app.route_document_data import (
    DocumentOptions,
    RouteDocumentData,
    RouteGeometryData,
    RouteSection,
)
from app.route_passport_docx import (
    DOCX_MIME,
    PROFILES,
    build_route_passport,
    passport_filename,
)


W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
WP = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"


def _data(*, forward=True, backward=True):
    def section(direction, rows):
        return RouteSection(direction, tuple(rows))

    return RouteDocumentData(
        route_id=1,
        route_number="42/А\r\n",
        route_name="Центральный маршрут",
        start_point="Вокзал",
        end_point="Аэропорт",
        version=2,
        forward=section("forward", [
            {"external_code": "F-01", "name": "Вокзал", "address": "Площадь, 1",
             "latitude": 56.123, "longitude": 35.987, "distance_from_prev_km": 0,
             "cumulative_km": 0},
            {"external_code": "F-02", "name": "Площадь", "address": None,
             "latitude": None, "longitude": None, "distance_from_prev_km": 2.5,
             "cumulative_km": 2.5},
            {"external_code": "F-03", "name": "Набережная", "address": "Речная, 7",
             "latitude": 56.5, "longitude": 36.2, "distance_from_prev_km": 4.25,
             "cumulative_km": 6.75},
        ] if forward else []),
        backward=section("backward", [
            {"external_code": "B-01", "name": "Аэропорт", "address": "Терминал, 1",
             "latitude": 56.456, "longitude": 36.123, "distance_from_prev_km": 0,
             "cumulative_km": 0},
        ] if backward else []),
        depot_out=section("depot_out", [{"external_code": "DEPOT", "name": "Депо"}]),
        depot_in=section("depot_in", [{"external_code": "DEPOT", "name": "Депо"}]),
        schedules={},
        geometries={
            "forward": RouteGeometryData("forward", "manual", 3, ((35.987, 56.123), (36.0, 56.2))),
            "backward": None,
        },
    )


def _options():
    return DocumentOptions("summer", "ЛЕТНИЙ ПЕРИОД", "ЛЕТО", date(2026, 8, 3))


def _offline_tile(url, timeout):
    raise OSError("offline")


def _xml(blob):
    with zipfile.ZipFile(io.BytesIO(blob)) as package:
        return ET.fromstring(package.read("word/document.xml")), package.namelist()


def _all_text(doc):
    return "\n".join(
        paragraph.text for paragraph in doc.paragraphs
    ) + "\n" + "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )


def test_builds_reopenable_portrait_package_with_required_content_and_maps():
    blob = build_route_passport(_data(), _options(), style="D", tile_loader=_offline_tile)

    assert DOCX_MIME == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    root, names = _xml(blob)
    assert root.tag == W + "document"
    assert any(name.startswith("word/media/") and name.endswith(".png") for name in names)
    doc = Document(io.BytesIO(blob))
    text = _all_text(doc)
    for expected in (
        "МИНИСТЕРСТВО ТРАНСПОРТА", "ПАСПОРТ МАРШРУТА РЕГУЛЯРНОГО СООБЩЕНИЯ",
        "СХЕМА ДВИЖЕНИЯ", "ОСТАНОВОЧНЫЕ ПУНКТЫ И РАССТОЯНИЯ",
        "ХАРАКТЕРИСТИКА АВТОМОБИЛЬНЫХ ДОРОГ",
        "F-01", "F-02", "B-01", "От предыдущего, км", "2.500", "________________", "Офлайн-схема",
    ):
        assert expected in text
    section = doc.sections[0]
    assert section.page_width < section.page_height
    assert section.page_width.cm == pytest.approx(21.0, abs=0.05)
    assert section.page_height.cm == pytest.approx(29.7, abs=0.05)
    assert (section.top_margin.cm, section.right_margin.cm, section.bottom_margin.cm, section.left_margin.cm) == pytest.approx((1.6, 1.6, 1.6, 1.6), abs=0.05)
    assert (section.header_distance.cm, section.footer_distance.cm) == pytest.approx((0.8, 0.8), abs=0.05)
    assert "Паспорт маршрута № 42/А" in section.header.paragraphs[0].text
    assert PROFILES[0].landscape is False


def test_filename_removes_path_and_linebreak_characters_and_uses_style_date():
    filename = passport_filename(_data(), _options(), "D")
    assert filename == "Паспорт_маршрута_42А_D_2026-08-03.docx"
    assert all(character not in filename for character in "\\/:*?\"<>|\r\n")


def test_manual_geometry_reaches_map_renderers(monkeypatch):
    from app import route_passport_docx as passport
    seen = []

    def scheme(section, geometry, **kwargs):
        seen.append(("scheme", section.direction, geometry.source if geometry else None))
        return passport.RenderedMap(_png(), False)

    def direction_map(section, geometry, **kwargs):
        seen.append(("map", section.direction, geometry.source if geometry else None))
        return passport.RenderedMap(_png(), False)

    monkeypatch.setattr(passport, "render_route_scheme", scheme)
    monkeypatch.setattr(passport, "render_direction_map", direction_map)
    build_route_passport(_data(backward=False), _options(), style="D", tile_loader=_offline_tile)
    assert ("scheme", "forward", "manual") in seen
    assert ("map", "forward", "manual") in seen


def _png():
    from PIL import Image
    image = Image.new("RGB", (32, 24), "white")
    result = io.BytesIO()
    image.save(result, format="PNG")
    return result.getvalue()


def test_tables_have_repeated_headers_and_fixed_consistent_dxa_geometry():
    blob = build_route_passport(_data(), _options(), style="D", tile_loader=_offline_tile)
    root, _ = _xml(blob)
    tables = root.findall(".//" + W + "tbl")
    assert tables
    for table in tables:
        width = table.find("./" + W + "tblPr/" + W + "tblW")
        grid = table.findall("./" + W + "tblGrid/" + W + "gridCol")
        cells = table.findall(".//" + W + "tcPr/" + W + "tcW")
        assert width is not None and width.get(W + "type") == "dxa"
        assert grid and cells
        expected = sum(int(column.get(W + "w")) for column in grid)
        assert int(width.get(W + "w")) == expected
        assert all(int(cell.get(W + "w")) in [int(c.get(W + "w")) for c in grid] for cell in cells)
    assert root.findall(".//" + W + "trPr/" + W + "tblHeader")


def test_footer_has_real_page_field_and_invalid_requests_fail():
    blob = build_route_passport(_data(), _options(), style="D", tile_loader=_offline_tile)
    with zipfile.ZipFile(io.BytesIO(blob)) as package:
        footer = ET.fromstring(package.read("word/footer1.xml"))
    field_types = [node.get(W + "fldCharType") for node in footer.findall(".//" + W + "fldChar")]
    instructions = [node.text.strip() for node in footer.findall(".//" + W + "instrText")]
    assert field_types == ["begin", "separate", "end"]
    assert instructions == ["PAGE"]
    with pytest.raises(ValueError) as unknown:
        build_route_passport(_data(), _options(), style="X")
    assert str(unknown.value) == "Оформление паспорта должно быть D или F"
    with pytest.raises(ValueError) as empty:
        build_route_passport(_data(forward=False, backward=False), _options(), style="D")
    assert str(empty.value) == "Для паспорта маршрута необходимо добавить остановки"


def test_filename_removes_all_control_characters_and_limits_route_token():
    filename = passport_filename(replace(_data(), route_number="42\x01\x1f" + "Ж" * 240), _options(), "D")
    route_token = filename.removeprefix("Паспорт_маршрута_").split("_D_", 1)[0]
    assert len(route_token) == 80
    assert all(ord(character) >= 32 for character in filename)


def _online_tile(url, timeout):
    from PIL import Image
    image = Image.new("RGBA", (256, 256), "white")
    output = io.BytesIO()
    image.save(output, format="PNG")
    return output.getvalue()


def test_map_captions_show_attribution_only_for_available_basemap():
    offline = _all_text(Document(io.BytesIO(build_route_passport(_data(), _options(), style="D", tile_loader=_offline_tile))))
    online = _all_text(Document(io.BytesIO(build_route_passport(_data(), _options(), style="D", tile_loader=_online_tile))))
    assert "Офлайн-схема" in offline
    assert "© OpenStreetMap contributors" not in offline
    assert "© OpenStreetMap contributors" in online
    assert "Офлайн-схема" not in online


def _table_with_header(doc, header):
    return next(table for table in doc.tables if table.rows[0].cells[0].text == "№" and header in [cell.text for cell in table.rows[0].cells])


def test_forward_stop_rows_preserve_sequence_and_distance_columns():
    doc = Document(io.BytesIO(build_route_passport(_data(), _options(), style="D", tile_loader=_offline_tile)))
    compact = _table_with_header(doc, "От предыдущего, км")
    detailed = _table_with_header(doc, "Расст. от пред., км")
    assert [cell.text for cell in compact.rows[0].cells] == ["№", "Код", "Остановочный пункт", "От предыдущего, км", "Нарастающим, км"]
    assert [cell.text for cell in compact.rows[1].cells] == ["1", "F-01", "Вокзал", "0.000", "0.000"]
    assert [cell.text for cell in compact.rows[2].cells] == ["2", "F-02", "Площадь", "2.500", "2.500"]
    assert [cell.text for cell in compact.rows[3].cells] == ["3", "F-03", "Набережная", "4.250", "6.750"]
    assert [cell.text for cell in detailed.rows[1].cells[:5]] == ["1", "F-01", "0.000", "0.000", "Вокзал"]
    assert [cell.text for cell in detailed.rows[2].cells[:5]] == ["2", "F-02", "2.500", "2.500", "Площадь"]
    assert [cell.text for cell in detailed.rows[3].cells[:5]] == ["3", "F-03", "4.250", "6.750", "Набережная"]


def _xml_table_with_header(root, header):
    return next(table for table in root.findall(".//" + W + "tbl") if header in "".join(table.itertext()))


@pytest.mark.parametrize(("header", "widths"), [
    ("От предыдущего, км", [500, 1300, 3691, 2300, 2300]),
    ("Муниципалитет / ОКАТО", [400, 800, 1150, 1050, 1650, 1650, 1050, 1050, 1291]),
])
def test_compact_and_detailed_tables_have_exact_fixed_word_geometry(header, widths):
    root, _ = _xml(build_route_passport(_data(), _options(), style="D", tile_loader=_offline_tile))
    table = _xml_table_with_header(root, header)
    properties = table.find("./" + W + "tblPr")
    layout = properties.find("./" + W + "tblLayout")
    width = properties.find("./" + W + "tblW")
    indent = properties.find("./" + W + "tblInd")
    margins = properties.find("./" + W + "tblCellMar")
    grid = table.findall("./" + W + "tblGrid/" + W + "gridCol")
    first_row = table.findall("./" + W + "tr")[0]
    first_widths = first_row.findall("./" + W + "tc/" + W + "tcPr/" + W + "tcW")
    assert layout.get(W + "type") == "fixed"
    assert (width.get(W + "type"), int(width.get(W + "w"))) == ("dxa", sum(widths))
    assert (indent.get(W + "type"), int(indent.get(W + "w"))) == ("dxa", 120)
    assert [int(node.get(W + "w")) for node in grid] == widths
    assert [int(node.get(W + "w")) for node in first_widths] == widths
    assert {name: int(margins.find("./" + W + name).get(W + "w")) for name in ("top", "bottom", "start", "end")} == {"top": 80, "bottom": 80, "start": 120, "end": 120}
    assert first_row.find("./" + W + "trPr/" + W + "tblHeader") is not None


def test_export_sanitizes_xml_illegal_controls_from_route_data():
    source = _data()
    forward = RouteSection("forward", tuple(
        {
            **stop,
            "external_code": "F\x00-01" if index == 0 else stop["external_code"],
            "name": f"{stop['name']}\x01",
            "address": "Речная\x02, 7",
            "municipality": "Город\x03",
        }
        for index, stop in enumerate(source.forward.stops)
    ))
    data = replace(
        source,
        route_number="42\x00",
        route_name="Центральный\x01 маршрут",
        start_point="Вокзал\x02",
        end_point="Аэропорт\x03",
        forward=forward,
    )
    blob = build_route_passport(data, _options(), style="D", tile_loader=_offline_tile)
    document = Document(io.BytesIO(blob))
    text = _all_text(document) + "\n" + document.sections[0].header.paragraphs[0].text
    assert all(ord(character) >= 32 or character in "\t\n\r" for character in text)


def test_landscape_f_uses_technical_cover_and_preserves_route_data():
    d_payload = build_route_passport(_data(), _options(), style="D", tile_loader=_offline_tile)
    f_payload = build_route_passport(_data(), _options(), style="F", tile_loader=_offline_tile)
    d_document = Document(io.BytesIO(d_payload))
    f_document = Document(io.BytesIO(f_payload))
    f_section = f_document.sections[0]
    f_text = _all_text(f_document)
    assert f_section.orientation == WD_ORIENT.LANDSCAPE
    assert f_section.page_width > f_section.page_height
    assert "ТЕХНИЧЕСКИЙ ПАСПОРТ МАРШРУТА" in f_text
    for value in ("42/А", "Центральный маршрут", "Длина прямого направления", "6.750 км", "Длина обратного направления", "0.000 км", "АВТОБУС", "03.08.2026"):
        assert value in f_text
    d_cells = [cell.text for table in d_document.tables for row in table.rows for cell in row.cells]
    f_cells = [cell.text for table in f_document.tables for row in table.rows for cell in row.cells]
    for value in ("Вокзал", "Площадь", "Набережная", "F-03", "4.250", "6.750"):
        assert value in d_cells
        assert value in f_cells
    d_root, _ = _xml(d_payload)
    f_root, _ = _xml(f_payload)
    d_table = _xml_table_with_header(d_root, "От предыдущего, км")
    f_table = _xml_table_with_header(f_root, "От предыдущего, км")
    d_width = int(d_table.find("./" + W + "tblPr/" + W + "tblW").get(W + "w"))
    f_width = int(f_table.find("./" + W + "tblPr/" + W + "tblW").get(W + "w"))
    f_grid = f_table.findall("./" + W + "tblGrid/" + W + "gridCol")
    assert f_width == sum(int(column.get(W + "w")) for column in f_grid)
    assert f_width > d_width
    assert passport_filename(_data(), _options(), "F").endswith("_F_2026-08-03.docx")


def test_landscape_f_image_extents_fit_the_page_composition_budget():
    payload = build_route_passport(_data(), _options(), style="F", tile_loader=_offline_tile)
    root, _ = _xml(payload)
    heights_cm = [int(extent.get("cy")) / 360000 for extent in root.findall(".//" + WP + "extent")]
    usable_height_cm = 21.0 - 2 * 1.4
    assert len(heights_cm) == 6
    assert sum(heights_cm[:2]) + 3.0 <= usable_height_cm
    assert all(height + 2.0 <= usable_height_cm for height in heights_cm[2:])


def test_landscape_f_technical_cover_marks_absent_direction():
    data = replace(_data(), backward=None)
    document = Document(io.BytesIO(build_route_passport(data, _options(), style="F", tile_loader=_offline_tile)))
    technical_cover = next(table for table in document.tables if "Параметр" in [cell.text for cell in table.rows[0].cells])
    values = {row.cells[0].text: row.cells[1].text for row in technical_cover.rows[1:]}
    assert values["Длина прямого направления"] == "6.750 км"
    assert values["Длина обратного направления"] == "—"
