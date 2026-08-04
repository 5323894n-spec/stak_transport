# -*- coding: utf-8 -*-
"""Deterministic fixture generation for manual render QA.

``generate_visual_fixtures`` writes ``passport-D.docx`` and ``passport-F.docx``
to an explicit output directory using an offline tile loader, so the result is
deterministic and safe to run without network access. The accompanying test
only ever writes inside ``tmp_path`` and never touches the repository.
"""

import datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document

from app.route_document_data import DocumentOptions, load_route_document_data
from app.route_passport_docx import build_route_passport
from tests.test_route_document_data import _database


def _offline_tiles(url, timeout):
    raise OSError("offline")


def generate_visual_fixtures(data, options, output_dir, *, tile_loader=_offline_tiles):
    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)
    written = []
    for style in ("D", "F"):
        payload = build_route_passport(
            data, options, style=style, tile_loader=tile_loader
        )
        path = output / f"passport-{style}.docx"
        path.write_bytes(payload)
        written.append(path)
    return written


def test_generate_visual_fixtures_writes_valid_packages(tmp_path):
    con, route_id = _database(tmp_path)
    try:
        data = load_route_document_data(con, route_id)
    finally:
        con.close()
    options = DocumentOptions(
        "winter", "ЗИМНИЙ ПЕРИОД", "ЗИМА", datetime.date(2026, 8, 3)
    )
    written = generate_visual_fixtures(data, options, tmp_path / "qa")
    assert [path.name for path in written] == [
        "passport-D.docx",
        "passport-F.docx",
    ]
    for path in written:
        with ZipFile(path) as package:
            names = package.namelist()
            assert "word/document.xml" in names
            assert any(name.startswith("word/media/") for name in names)
        document = Document(str(path))
        assert document.tables, "passport must contain tables"
        assert document.paragraphs, "passport must contain text"
