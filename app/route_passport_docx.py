# -*- coding: utf-8 -*-
"""Editable portrait Word passport for a regular public transport route."""
from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
import re

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .route_passport_maps import RenderedMap, render_direction_map, render_route_scheme

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_ACCENT, _DARK, _FILL = "17365D", "10243E", "E8EEF5"


@dataclass(frozen=True)
class PassportProfile:
    code: str
    landscape: bool
    accent: str
    page_width_cm: float
    page_height_cm: float
    margin_cm: float
    map_width_cm: float


PROFILES = (
    PassportProfile("D", False, _ACCENT, 21.0, 29.7, 1.6, 17.8),
    PassportProfile("F", True, "276678", 29.7, 21.0, 1.4, 26.9),
)
_BY_CODE = {profile.code: profile for profile in PROFILES}


def _profile(style):
    try:
        return _BY_CODE[style]
    except KeyError:
        raise ValueError("Оформление паспорта должно быть D или F") from None


def _display_text(value):
    """Return XML 1.0-safe text while retaining all legal Unicode and whitespace."""
    if value is None:
        return ""
    return "".join(
        character for character in str(value)
        if character in "\t\n\r" or 0x20 <= ord(character) <= 0xD7FF
        or 0xE000 <= ord(character) <= 0xFFFD or 0x10000 <= ord(character) <= 0x10FFFF
    )

def passport_filename(data, options, style):
    _profile(style)
    route = re.sub(r'[\\/:*?"<>|\x00-\x1f\x7f]+', "", _display_text(getattr(data, "route_number", ""))).strip()[:80] or "без_номера"
    return f"Паспорт_маршрута_{route}_{style}_{options.effective_date.isoformat()}.docx"


def _font(font, size=None, bold=None, color=None):
    font.name = "Arial"
    for key in ("ascii", "hAnsi", "cs"):
        font._element.rPr.rFonts.set(qn(f"w:{key}"), "Arial")
    if size is not None: font.size = Pt(size)
    if bold is not None: font.bold = bold
    if color: font.color.rgb = RGBColor.from_string(color)


def _section(section, profile):
    section.orientation = WD_ORIENT.LANDSCAPE if profile.landscape else WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = Cm(profile.page_width_cm), Cm(profile.page_height_cm)
    section.top_margin = section.bottom_margin = Cm(profile.margin_cm)
    section.left_margin = section.right_margin = Cm(profile.margin_cm)
    section.header_distance = section.footer_distance = Cm(0.8)


def _styles(doc, profile):
    normal = doc.styles["Normal"]; _font(normal.font, 10.5)
    normal.paragraph_format.space_after, normal.paragraph_format.line_spacing = Pt(4), 1.15
    title = doc.styles["Title"]; _font(title.font, 20, True, profile.accent)
    title.paragraph_format.alignment, title.paragraph_format.space_after = WD_ALIGN_PARAGRAPH.CENTER, Pt(6)
    subtitle = doc.styles["Subtitle"]; _font(subtitle.font, 10.5, color=_DARK)
    subtitle.paragraph_format.space_after = Pt(4)
    for name, size, color, before, after in (("Heading 1", 15, profile.accent, 10, 6), ("Heading 2", 12, profile.accent, 8, 4), ("Heading 3", 10.5, _DARK, 6, 3)):
        style = doc.styles[name]; _font(style.font, size, True, color)
        style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(before), Pt(after)
        style.paragraph_format.line_spacing = 1.15


def _page_field(paragraph):
    run = OxmlElement("w:r")
    for tag, value in (("fldChar", "begin"), ("instrText", " PAGE "), ("fldChar", "separate"), ("t", "1"), ("fldChar", "end")):
        item = OxmlElement(f"w:{tag}")
        if tag == "fldChar": item.set(qn("w:fldCharType"), value)
        else: item.text = value
        run.append(item)
    paragraph._p.append(run)


def _furniture(section, number):
    header = section.header.paragraphs[0]; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.add_run(f"Паспорт маршрута № {_display_text(number) or '________________'}")
    footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Лист "); _page_field(footer)


def _shading(cell):
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), _FILL); shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def _geometry(table, widths):
    total_width = sum(widths)
    table.autofit = False; pr = table._tbl.tblPr
    layout = pr.first_child_found_in("w:tblLayout")
    if layout is None: layout = OxmlElement("w:tblLayout"); pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for name, width in (("tblW", total_width), ("tblInd", 120)):
        node = pr.find(qn(f"w:{name}"))
        if node is None: node = OxmlElement(f"w:{name}"); pr.append(node)
        node.set(qn("w:w"), str(width)); node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for node in list(grid): grid.remove(node)
    for width in widths:
        node = OxmlElement("w:gridCol"); node.set(qn("w:w"), str(width)); grid.append(node)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            node = cell._tc.get_or_add_tcPr().find(qn("w:tcW")); node.set(qn("w:w"), str(width)); node.set(qn("w:type"), "dxa")
    margins = OxmlElement("w:tblCellMar")
    for name, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = OxmlElement(f"w:{name}"); node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa"); margins.append(node)
    pr.append(margins)
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{name}"); node.set(qn("w:val"), "single"); node.set(qn("w:sz"), "4"); node.set(qn("w:color"), "AAB7C4"); borders.append(node)
    pr.append(borders)



def _table_widths(profile, widths):
    target = round((profile.page_width_cm - 2 * profile.margin_cm) / 2.54 * 1440)
    scaled = [round(width * target / sum(widths)) for width in widths[:-1]]
    return tuple((*scaled, target - sum(scaled)))

def _table(doc, headers, rows, widths, numeric=()):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]; cell.text = value; _shading(cell)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs: _font(run.font, 9, True, _DARK)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]; cell.text = _display_text(value)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if index in numeric else WD_ALIGN_PARAGRAPH.LEFT
            for run in cell.paragraphs[0].runs: _font(run.font, 9.2)
    _geometry(table, widths)
    return table


def _heading(doc, text, level=1): return doc.add_paragraph(text, style=f"Heading {level}")
def _label(section): return "Прямое направление" if section.direction == "forward" else "Обратное направление"
def _route(data):
    route_name = _display_text(getattr(data, "route_name", ""))
    endpoints = tuple(_display_text(value) for value in (getattr(data, "start_point", ""), getattr(data, "end_point", "")))
    return route_name or " - ".join(value for value in endpoints if value) or "________________"

def _num(value):
    try: return f"{float(value):.3f}" if value not in (None, "") else "—"
    except (TypeError, ValueError): return "—"


def _picture(doc, rendered, width, caption=False):
    doc.add_picture(BytesIO(rendered.png), width=Cm(width))
    if caption:
        text = f"Картографическая основа: {rendered.attribution}" if rendered.basemap_available else "Офлайн-схема: картографическая подложка недоступна"
        doc.add_paragraph(text, style="Subtitle")



def _direction_length(section):
    if section is None or not getattr(section, "stops", ()):
        return "—"

    try:
        return f"{sum(float(row.get('distance_from_prev_km') or 0) for row in section.stops):.3f} км"
    except (TypeError, ValueError):
        return "—"


def _technical_cover(doc, data, options, profile):
    doc.add_paragraph("ТЕХНИЧЕСКИЙ ПАСПОРТ МАРШРУТА", style="Title")
    _table(doc, ("Параметр", "Значение"), (
        ("Номер маршрута", _display_text(getattr(data, "route_number", "")) or "________________"),
        ("Наименование маршрута", _route(data)),
        ("Длина прямого направления", _direction_length(data.forward)),
        ("Длина обратного направления", _direction_length(data.backward)),
        ("Вид транспорта", "АВТОБУС"),
        ("Дата состояния", options.effective_date.strftime("%d.%m.%Y")),
    ), _table_widths(profile, (3600, 6491)))

def _cover(doc, data, options, profile):
    if profile.code == "F":
        return _technical_cover(doc, data, options, profile)
    paragraph = doc.add_paragraph(); paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before, paragraph.paragraph_format.space_after = Pt(54), Pt(34)
    run = paragraph.add_run("МИНИСТЕРСТВО ТРАНСПОРТА"); _font(run.font, 11, True, _ACCENT)
    doc.add_paragraph("ПАСПОРТ МАРШРУТА РЕГУЛЯРНОГО СООБЩЕНИЯ", style="Title")
    paragraph = doc.add_paragraph(); paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"№ {_display_text(getattr(data, 'route_number', '')) or '________________'}"); _font(run.font, 16, True, _DARK)
    paragraph = doc.add_paragraph(); paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(_route(data)); _font(run.font, 14, True)
    for text in ("Реестровый номер: ________________________________", "Вид транспорта: АВТОБУС", "Организация перевозчик: ________________________________"):
        paragraph = doc.add_paragraph(text); paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _approval(doc, data, options, profile):
    _heading(doc, "УТВЕРЖДЕНИЕ И РЕГИСТРАЦИЯ")
    for text in ("УТВЕРЖДАЮ: ________________________________", "Должность: ____________________________________________", "Подпись: ____________________   Дата: ________________"):
        doc.add_paragraph(text)
    _table(doc, ("Реквизит", "Значение"), (("Номер маршрута", _display_text(getattr(data, "route_number", "")) or "________________"), ("Наименование", _route(data)), ("Дата состояния", options.effective_date.strftime("%d.%m.%Y")), ("Вид транспорта", "АВТОБУС"), ("Вид маршрута", "________________"), ("ОКАТО", "________________")), _table_widths(profile, (3000, 7091)))
    doc.add_paragraph("Регистрационный номер: ________________________________")


def _render_size(profile, kind):
    if profile.code == "F":
        return (1600, 400) if kind == "full" else (1600, 700)
    return (1200, 620) if kind == "full" else (1200, 800)


def _full_scheme(doc, data, profile, directions):
    _heading(doc, "СХЕМА ДВИЖЕНИЯ МАРШРУТА")
    for section in directions:
        _heading(doc, _label(section), 3)
        _picture(doc, render_route_scheme(section, data.geometries.get(section.direction), size=_render_size(profile, "full")), profile.map_width_cm)


def _direction_scheme(doc, data, profile, section):
    _heading(doc, f"СХЕМА ДВИЖЕНИЯ - {_label(section).upper()}")
    _picture(doc, render_route_scheme(section, data.geometries.get(section.direction), size=_render_size(profile, "direction")), profile.map_width_cm)


def _direction_map(doc, data, profile, section, tile_loader):
    _heading(doc, f"КАРТА МАРШРУТА - {_label(section).upper()}")
    _picture(doc, render_direction_map(section, data.geometries.get(section.direction), size=_render_size(profile, "direction"), tile_loader=tile_loader), profile.map_width_cm, True)


def _compact(doc, directions, profile):
    _heading(doc, "ОСТАНОВОЧНЫЕ ПУНКТЫ И РАССТОЯНИЯ")
    for section in directions:
        _heading(doc, _label(section), 2)
        _table(doc, ("№", "Код", "Остановочный пункт", "От предыдущего, км", "Нарастающим, км"), ((number, row.get("external_code") or "—", row.get("name") or "—", _num(row.get("distance_from_prev_km")), _num(row.get("cumulative_km"))) for number, row in enumerate(section.stops, 1)), _table_widths(profile, (500, 1300, 3691, 2300, 2300)), (0, 3, 4))


def _detailed(doc, section, profile):
    _heading(doc, f"ПОДРОБНЫЕ ОСТАНОВОЧНЫЕ ПУНКТЫ - {_label(section).upper()}")
    _table(doc, ("№", "Код", "Расст. от пред., км", "Нарастающим, км", "Остановка", "Адрес", "Широта", "Долгота", "Муниципалитет / ОКАТО"), ((number, row.get("external_code") or "—", _num(row.get("distance_from_prev_km")), _num(row.get("cumulative_km")), row.get("name") or "—", row.get("address") or "—", _num(row.get("latitude")), _num(row.get("longitude")), row.get("municipality") or row.get("okato") or "________________") for number, row in enumerate(section.stops, 1)), _table_widths(profile, (400, 800, 1150, 1050, 1650, 1650, 1050, 1050, 1291)), (0, 2, 3, 6, 7))


def _roads(doc):
    _heading(doc, "ХАРАКТЕРИСТИКА АВТОМОБИЛЬНЫХ ДОРОГ")
    for text in ("Улицы (дороги), категория: _____________________________________________", "Ширина проезжей части, покрытие: ______________________________________", "Пересечения, мосты, железнодорожные переезды: _________________________", "Опасные участки и меры безопасности: __________________________________", "Примечания: ___________________________________________________________"):
        doc.add_paragraph(text)


def build_route_passport(data, options, *, style, tile_loader=None) -> bytes:
    """Build the approved D profile; depot legs are intentionally excluded."""
    profile = _profile(style)
    directions = tuple(item for item in (getattr(data, "forward", None), getattr(data, "backward", None)) if item is not None and getattr(item, "stops", ()))
    if not directions: raise ValueError("Для паспорта маршрута необходимо добавить остановки")
    doc = Document(); section = doc.sections[0]
    _section(section, profile); _styles(doc, profile); _furniture(section, getattr(data, "route_number", "")); _cover(doc, data, options, profile)
    doc.add_page_break(); _approval(doc, data, options, profile)
    doc.add_page_break(); _full_scheme(doc, data, profile, directions)
    for item in directions:
        doc.add_page_break(); _direction_scheme(doc, data, profile, item)
        doc.add_page_break(); _direction_map(doc, data, profile, item, tile_loader)
    doc.add_page_break(); _compact(doc, directions, profile)
    for item in directions: doc.add_page_break(); _detailed(doc, item, profile)
    doc.add_page_break(); _roads(doc)
    output = BytesIO(); doc.save(output); return output.getvalue()
